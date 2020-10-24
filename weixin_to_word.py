import requests
from pyquery import PyQuery as pq
from docx import Document
import time
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Inches
import re
from io import BytesIO
from PIL import Image


def add_pic(picurl, picwidth):
    # 获取图片二进制数据并保存
    picresponse = requests.get(picurl, headers=headers)
    picpath = './pic/' + time.strftime('%Y-%m-%d %H') + '.png'
    try:
        # 该方法保存的部分图片（疑似来自ios端）无法用python_docx写入docx
        with open(picpath, 'wb') as f:
            f.write(picresponse.content)
        paragraph = doc.add_paragraph()
        # 图片居中设置
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = paragraph.add_run("")
        run.add_picture(picpath, width=Inches(5 / 600 * picwidth))
    except:
        img = Image.open(BytesIO(picresponse.content))
        img.save(picpath)
        paragraph = doc.add_paragraph()
        # 图片居中设置
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = paragraph.add_run("")
        run.add_picture(picpath, width=Inches(5 / 600 * picwidth))


start = time.perf_counter()
doc = Document()
url = 'https://mp.weixin.qq.com/s?__biz=MzUxMDMzODg2Ng==&mid=100000279&idx=1&sn=686a9f916f1b91e600965ce66e53ea64&chksm=7905304a4e72b95c9d82f096a94e12cb243fd8495d86c9e7b86e1d796e665d85f7d84a4f8739&mpshare=1&scene=1&srcid=1022NjHRw4T7N45fmKrJ9XaP&sharer_sharetime=1603336476778&sharer_shareid=dc642653ff36899e41758c9327365611&key=25b7ee6511d12c932a17199006e915fa6915fccbb65c6152ce8b00de41b7beed5b93fcc1d53e0c7469d0a50b67a68351ef4d415b7d68ebcdff4e55a59af0f039785f648293c2d831b7734531dd2f3291ae08e63f763c6a2e5c6a1648adcc898932904a7deb247aad10891507f32091a2350c6a3a9eb85c49e374b0dd56ebbf66&ascene=1&uin=MzY1MjAxMjE3OA%3D%3D&devicetype=Windows+10+x64&version=62090529&lang=zh_CN&exportkey=A%2BiLnA7AOCTfUsHDPr86THE%3D&pass_ticket=QT3EEZ0hENgySf%2B9EqvjGk0caDFqBWrReGTHWrBCdx2kRh2XMZQ6LpqTcNG9nlOm&wx_header=0'
headers = {
    'user-agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/86.0.4240.75 Safari/537.36'
}
response = requests.get(url, headers=headers)
htmldoc = pq(response.text)
box = htmldoc('#js_content')
head1box = box('[style="white-space: normal;line-height: 1.75em;"]').items()
head1 = [x.text() for x in head1box if x.text()]
head21 = ['写作理论', '文章结构', '句型模板', ' 细说表征', '论文投稿', '写作细节']
head22 = ['Origin基础教程', 'Origin插件', 'Origin常见问题汇总', '平面作图软件', 'PPT&Excel', '其他作图相关软件',
          '备注：三维作图，请大家关注公众号“3D科研绘图”，部分内容汇总如下：']
head23 = ['论文检索', '文献管理', '其他']
head2 = head21 + head22 + head23
print(head2)
head3box = box('[data-tools="新媒体排版"]').items()
picformer = 'https://mmbiz'
head2num = 0
artcount = 0
for i in head3box:
    ullist = i('ul').items()
    for ul in ullist:
        if head2num < 6:
            if head2num == 0:
                doc.add_heading(head1[0], level=1)
                doc.add_heading(head2[head2num], level=2)
            else:
                doc.add_heading(head2[head2num], level=2)
        elif head2num < 13:
            if head2num == 6:
                doc.add_heading(head1[1], level=1)
                doc.add_heading(head2[head2num], level=2)
            else:
                doc.add_heading(head2[head2num], level=2)
        elif head2num < 16:
            if head2num == 13:
                doc.add_heading(head1[1], level=1)
                doc.add_heading(head2[head2num], level=2)
            else:
                doc.add_heading(head2[head2num], level=2)
        head2num = head2num + 1
        alist = ul('a').items()
        for a in alist:
            arturl = a.attr('href')
            head3 = a.text()
            doc.add_heading(head3, level=3)
            print(head3)
            artresponse = requests.get(arturl, headers=headers)

            # 对于公众号已迁移的重新定位原文章
            if '该公众号已迁移' in artresponse.text:
                arturl = re.findall("transferTargetLink = '(\S+)'", artresponse.text)[0]
                artresponse = requests.get(arturl, headers=headers)
            artdoc = pq(artresponse.text)
            artcontent = artdoc('#js_content').children().items()
            for parai in artcontent:
                if not parai.html():
                    continue
                elif ('135&#x7F16;&#x8F91;&#x5668' in parai.html()) or ('Powered by 135editor.com' in parai.html()):
                    bjqdoc = pq(parai.html())
                    plist = bjqdoc('p').items()
                    for p in plist:
                        if not p.html():
                            continue
                        elif re.findall('src="https://mmbiz(\S+?)"', p.html()) and ('img' in p.html()):
                            picurl = picformer + re.findall('src="https://mmbiz(\S+?)"', p.html())[0]
                            picwidth = re.findall('width="(\d+)\S*?"', p.html())
                            if picwidth:
                                picwidth = eval(picwidth[0])
                            else:
                                picwidth = 478
                            print(picurl, '\n', picwidth)
                            add_pic(picurl, picwidth)
                        doc.add_paragraph(p.text())
                elif re.findall('src="https://mmbiz(\S+?)"', parai.html()) and ('img' in parai.html()):
                    picurl = picformer + re.findall('src="https://mmbiz(\S+?)"', parai.html())[0]
                    picwidth = re.findall('width="(\d+)\S*?"', parai.html())
                    if picwidth:
                        picwidth = eval(picwidth[0])
                    else:
                        picwidth = 478
                    print(picurl, '\n', picwidth)
                    try:
                        add_pic(picurl, picwidth)
                    except:
                        print('Error')
                doc.add_paragraph(parai.text())
            doc.add_paragraph('原文链接：' + arturl)
            artcount = artcount + 1
            print('已完成{}篇文章'.format(artcount))
doc.save('科研技巧专栏汇总2.0.docx')
end = time.perf_counter()
print(end - start)